from fpdf import FPDF
from docx import Document
from datetime import datetime
import os


class ReportPDF(FPDF):
    def header(self):
        self.set_font("Helvetica", "B", 14)
        self.set_text_color(26, 58, 110)
        self.cell(0, 10, "AI Email Content Checker - Tata Steel", ln=True, align="C")
        self.set_font("Helvetica", "", 9)
        self.set_text_color(120, 120, 120)
        self.cell(0, 6, f"Generated: {datetime.now().strftime('%d %B %Y, %I:%M %p')}",
                  ln=True, align="C")
        self.ln(4)

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")


def export_as_pdf(subject, recipient, results, overall, corrections, links_data):
    """Generate a professional PDF report."""
    os.makedirs("reports", exist_ok=True)
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    filename = f"reports/report_{timestamp}.pdf"

    pdf = ReportPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Email details section
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 8, "Email Details", ln=True)
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 6, f"Subject   : {subject}", ln=True)
    pdf.cell(0, 6, f"Recipient : {recipient}", ln=True)
    pdf.cell(0, 6, f"Overall Score: {overall}/100", ln=True)
    pdf.ln(4)

    # Agent results
    for result in results:
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(26, 58, 110)
        # Clean emoji from agent name for PDF
        agent_name = result['agent'].encode('latin-1', 'ignore').decode('latin-1')
        pdf.cell(0, 8, agent_name, ln=True)
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(0, 0, 0)
        # Clean text for PDF encoding
        raw_text = result['raw'].encode('latin-1', 'ignore').decode('latin-1')
        pdf.multi_cell(0, 5, raw_text)
        pdf.ln(3)

    # Grammar corrections table
    if corrections:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(26, 58, 110)
        pdf.cell(0, 8, "Grammar & Style Corrections", ln=True)
        pdf.ln(2)

        # Table headers
        pdf.set_fill_color(230, 236, 245)
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(60, 7, "ORIGINAL", border=1, fill=True)
        pdf.cell(60, 7, "CORRECTION", border=1, fill=True)
        pdf.cell(70, 7, "EXPLANATION", border=1, fill=True, ln=True)

        # Table rows
        pdf.set_font("Helvetica", "", 8)
        for c in corrections:
            orig = str(c.get("original", ""))[:80].encode('latin-1', 'ignore').decode('latin-1')
            corr = str(c.get("correction", ""))[:80].encode('latin-1', 'ignore').decode('latin-1')
            expl = str(c.get("explanation", ""))[:100].encode('latin-1', 'ignore').decode('latin-1')
            h = 12
            pdf.multi_cell(60, h, orig, border=1)
            # Reset position for next cells
            pdf.set_xy(pdf.get_x() + 60, pdf.get_y() - h)
            pdf.multi_cell(60, h, corr, border=1)
            pdf.set_xy(pdf.get_x() + 120, pdf.get_y() - h)
            pdf.multi_cell(70, h, expl, border=1)

    # Links section
    if links_data:
        pdf.ln(4)
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(26, 58, 110)
        pdf.cell(0, 8, "Hyperlink Verification", ln=True)
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(0, 0, 0)
        for link in links_data:
            status = "✓ Working" if link["working"] else "✗ Broken"
            line = f"{status} | {link['url']} | Status: {link['status_code']}"
            line = line.encode('latin-1', 'ignore').decode('latin-1')
            pdf.cell(0, 6, line, ln=True)

    pdf.output(filename)
    return filename


def export_as_docx(subject, recipient, results, overall, corrections, links_data):
    """Generate a Word document report."""
    os.makedirs("reports", exist_ok=True)
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    filename = f"reports/report_{timestamp}.docx"

    doc = Document()

    # Title
    doc.add_heading("AI Email Content Checker — Tata Steel", 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y, %I:%M %p')}")

    # Email details
    doc.add_heading("Email Details", 1)
    doc.add_paragraph(f"Subject: {subject}")
    doc.add_paragraph(f"Recipient: {recipient}")
    doc.add_paragraph(f"Overall Score: {overall}/100")

    # Agent results
    doc.add_heading("Agent Analysis Results", 1)
    for result in results:
        doc.add_heading(result["agent"], 2)
        doc.add_paragraph(result["raw"])

    # Grammar corrections table
    if corrections:
        doc.add_heading("Grammar & Style Corrections", 1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "ORIGINAL"
        hdr[1].text = "CORRECTION"
        hdr[2].text = "EXPLANATION"
        for c in corrections:
            row = table.add_row().cells
            row[0].text = str(c.get("original", ""))
            row[1].text = str(c.get("correction", ""))
            row[2].text = str(c.get("explanation", ""))

    # Links
    if links_data:
        doc.add_heading("Hyperlink Verification", 1)
        for link in links_data:
            status = "✓ Working" if link["working"] else "✗ Broken"
            doc.add_paragraph(
                f"{status} — {link['url']} (HTTP {link['status_code']})"
            )

    doc.save(filename)
    return filename